# -*- coding: utf-8 -*-
"""Assemble the SentriZK JTEC review paper (.docx) from workflow JSON outputs."""
import json, os, re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BUILD = r"C:\FYP_BCSS\SentriZK-InternalChat\Doc\Academic\Papers\_build"
OUT   = os.path.join(BUILD, "SentriZK_JTEC_Paper.docx")
FIG1  = os.path.join(BUILD, "fig1_architecture.png")

# ----------------------------- EDIT THESE -----------------------------
TITLE = ("SentriZK: A Review of Privacy-Preserving Secure Messaging via "
         "Zero-Knowledge Authentication, Insider Threat Detection, and On-Device AI")
AUTHORS = "[First Author]¹, [Second Author]¹, [Third Author]¹, [Supervisor Name]²"
AFFIL = ("¹Bachelor of Computer Science (Cyber Security), [Faculty], [University], [City], Malaysia\n"
         "²[Department], [University], [City], Malaysia\n"
         "Email: [corresponding.author]@[university].edu.my")
# ----------------------------------------------------------------------

SECTION_ORDER = ["sec_intro","sec_method","sec_d1","sec_d2","sec_d3","sec_synth","sec_arch","sec_concl"]
PAPER_IDS = [f"P{i}" for i in range(1,13)]

def roman(n):
    vals=[(1000,'M'),(900,'CM'),(500,'D'),(400,'CD'),(100,'C'),(90,'XC'),(50,'L'),
          (40,'XL'),(10,'X'),(9,'IX'),(5,'V'),(4,'IV'),(1,'I')]
    r=''
    for v,s in vals:
        while n>=v: r+=s; n-=v
    return r

def load_json(name):
    p=os.path.join(BUILD,name)
    with open(p,'r',encoding='utf-8') as f:
        return json.load(f)

# ---- load digests (references) ----
digs={}
for pid in PAPER_IDS:
    fp=os.path.join(BUILD,f"dig_{pid}.json")
    if os.path.exists(fp):
        try: digs[pid]=load_json(f"dig_{pid}.json")
        except Exception as e: print("WARN bad digest",pid,e)
print(f"Loaded {len(digs)} digests")

# ---- load sections ----
sections=[]
for name in SECTION_ORDER:
    sections.append((name, load_json(name+".json")))

abstract_data = load_json("abstract.json")

# ---- citation numbering by first appearance ----
TOK = re.compile(r"\bP(\d+)\b")
order=[]
def scan(text):
    for m in TOK.finditer(text or ""):
        pid="P"+m.group(1)
        if pid not in order and pid in PAPER_IDS:
            order.append(pid)
for _,sec in sections:
    for b in sec.get("blocks",[]):
        scan(b.get("text",""))
    for t in sec.get("tables",[]):
        scan(t.get("caption",""))
        for row in t.get("rows",[]):
            for c in row: scan(c)
# any cited-but-missing? any uncited?
cited=set(order)
uncited=[p for p in PAPER_IDS if p in digs and p not in cited]
for p in uncited:
    order.append(p)  # append uncited at the end so all refs appear
num={pid:i+1 for i,pid in enumerate(order)}
print("Citation order:", order)
if uncited: print("WARNING uncited (appended):", uncited)
missing_dig=[p for p in order if p not in digs]
if missing_dig: print("WARNING cited but no digest/reference:", missing_dig)

def renumber(text):
    def rep(m):
        pid="P"+m.group(1)
        return str(num.get(pid, m.group(0)))
    return TOK.sub(rep, text or "")

# =========================== DOCX BUILD ===========================
doc=Document()

# base style
normal=doc.styles['Normal']
normal.font.name='Times New Roman'
normal.font.size=Pt(10)
rpr=normal.element.get_or_add_rPr(); rfonts=rpr.get_or_add_rFonts()
rfonts.set(qn('w:ascii'),'Times New Roman'); rfonts.set(qn('w:hAnsi'),'Times New Roman')

def set_cols(section,n,space=425):
    sectPr=section._sectPr
    cols=sectPr.find(qn('w:cols'))
    if cols is None:
        cols=OxmlElement('w:cols'); sectPr.append(cols)
    cols.set(qn('w:num'),str(n)); cols.set(qn('w:space'),str(space)); cols.set(qn('w:equalWidth'),'1')

def page_setup(section):
    section.page_height=Inches(11.69); section.page_width=Inches(8.27)  # A4
    section.top_margin=Inches(0.7); section.bottom_margin=Inches(0.7)
    section.left_margin=Inches(0.6); section.right_margin=Inches(0.6)

col_plan=[1]  # gap0 = title/abstract block (full width)

def break_to(n):
    doc.add_section(WD_SECTION.CONTINUOUS); col_plan.append(n)

def para(text,size=10,bold=False,italic=False,align=WD_ALIGN_PARAGRAPH.JUSTIFY,
         indent=0.0,before=0,after=2,line=1.0,color=None,name='Times New Roman'):
    p=doc.add_paragraph(); p.alignment=align
    pf=p.paragraph_format; pf.space_before=Pt(before); pf.space_after=Pt(after)
    pf.line_spacing=line
    if indent: pf.first_line_indent=Inches(indent)
    r=p.add_run(text); r.font.size=Pt(size); r.bold=bold; r.italic=italic
    r.font.name=name
    if color: r.font.color.rgb=color
    return p

def runs_para(segments,size=10,align=WD_ALIGN_PARAGRAPH.JUSTIFY,indent=0.0,before=0,after=2,line=1.0):
    p=doc.add_paragraph(); p.alignment=align
    pf=p.paragraph_format; pf.space_before=Pt(before); pf.space_after=Pt(after); pf.line_spacing=line
    if indent: pf.first_line_indent=Inches(indent)
    for txt,b,i in segments:
        r=p.add_run(txt); r.font.size=Pt(size); r.bold=b; r.italic=i; r.font.name='Times New Roman'
    return p

# ---------- TITLE BLOCK (full width) ----------
para(TITLE,size=18,bold=True,align=WD_ALIGN_PARAGRAPH.CENTER,after=8,line=1.05)
para(AUTHORS,size=11,align=WD_ALIGN_PARAGRAPH.CENTER,after=2)
for line in AFFIL.split("\n"):
    para(line,size=9,italic=True,align=WD_ALIGN_PARAGRAPH.CENTER,after=1)

# Abstract
abs_text=abstract_data.get("abstract","")
runs_para([("Abstract—",True,True),(abs_text,True,False)],size=9,before=8,after=3)
# Index terms
terms=abstract_data.get("index_terms",[])
terms_str="; ".join(t.strip().rstrip('.') for t in terms)+"."
runs_para([("Index Terms—",True,True),(terms_str,False,True)],size=9,after=4)

# switch to two columns for the body
break_to(2)

# ---------- BODY SECTIONS ----------
def add_heading(text,n):
    p=para(f"{roman(n)}. {text.upper()}",size=10,bold=True,
           align=WD_ALIGN_PARAGRAPH.CENTER,before=10,after=4)
    return p

def add_table(tbl):
    cols=tbl.get("columns",[]); rows=tbl.get("rows",[])
    cap=renumber(tbl.get("caption",""))
    # full width: break out to 1 col
    break_to(1)
    para(cap,size=8,bold=True,align=WD_ALIGN_PARAGRAPH.CENTER,before=6,after=3)
    if cols:
        t=doc.add_table(rows=1,cols=len(cols)); t.style='Table Grid'
        t.alignment=WD_ALIGN_PARAGRAPH.CENTER
        hdr=t.rows[0].cells
        for j,c in enumerate(cols):
            hdr[j].text=''
            pp=hdr[j].paragraphs[0]; pp.alignment=WD_ALIGN_PARAGRAPH.CENTER
            rr=pp.add_run(str(c)); rr.bold=True; rr.font.size=Pt(8); rr.font.name='Times New Roman'
        for row in rows:
            cells=t.add_row().cells
            for j in range(len(cols)):
                val=row[j] if j<len(row) else ''
                cells[j].text=''
                pp=cells[j].paragraphs[0]; pp.alignment=WD_ALIGN_PARAGRAPH.LEFT
                rr=pp.add_run(renumber(str(val))); rr.font.size=Pt(8); rr.font.name='Times New Roman'
    # back to two columns
    break_to(2)

sec_titles={"sec_intro":"INTRODUCTION","sec_method":"REVIEW METHODOLOGY"}
n=0
for name,sec in sections:
    n+=1
    title=sec.get("section_title") or sec_titles.get(name,name)
    add_heading(title,n)
    for b in sec.get("blocks",[]):
        typ=b.get("type","body"); txt=renumber(b.get("text",""))
        if typ=="subheading":
            para(txt,size=10,italic=True,align=WD_ALIGN_PARAGRAPH.LEFT,before=6,after=2)
        else:
            para(txt,size=10,indent=0.2,after=3,line=1.0)
            # insert Fig 1 right after the block that references it (architecture section)
            if name=="sec_arch" and ("Fig. 1" in b.get("text","") or "Fig.1" in b.get("text","")):
                if os.path.exists(FIG1):
                    fp=doc.add_paragraph(); fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
                    fp.add_run().add_picture(FIG1,width=Inches(3.2))
                    para("Fig. 1. Conceptual layered architecture of SentriZK.",
                         size=8,align=WD_ALIGN_PARAGRAPH.CENTER,before=2,after=6)
    for t in sec.get("tables",[]):
        add_table(t)

# ---------- REFERENCES ----------
add_heading("References",n+1) if False else None
# References heading without number (IEEE convention)
para("REFERENCES",size=10,bold=True,align=WD_ALIGN_PARAGRAPH.CENTER,before=10,after=4)
for i,pid in enumerate(order):
    ref=digs.get(pid,{}).get("ieee_reference","").strip()
    if not ref:
        ref=f"[reference missing for {pid}]"
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    pf=p.paragraph_format; pf.space_after=Pt(2); pf.line_spacing=1.0
    pf.left_indent=Inches(0.22); pf.first_line_indent=Inches(-0.22)
    r=p.add_run(f"[{i+1}]\t{ref}"); r.font.size=Pt(8); r.font.name='Times New Roman'

# ---------- finalize columns + page setup on every section ----------
secs=doc.sections
for g,s in enumerate(secs):
    page_setup(s)
    set_cols(s, col_plan[g] if g<len(col_plan) else 2)

doc.save(OUT)

# stats
def count_words():
    w=0
    for _,sec in sections:
        for b in sec.get("blocks",[]): w+=len(b.get("text","").split())
    w+=len(abs_text.split())
    return w
print("SAVED:",OUT)
print("Sections:",len(sections)," References:",len(order)," Approx body words:",count_words())
